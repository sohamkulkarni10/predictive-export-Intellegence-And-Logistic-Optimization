"""Convert project markdown documentation to a Word .docx file."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\Lenovo\Desktop\Export_AI")
MD = ROOT / "ExportIntel_AI_Complete_Project_Documentation.md"
OUT = ROOT / "ExportIntel_AI_Complete_Project_Documentation.docx"


def add_runs_with_code(paragraph, line: str) -> None:
    i = 0
    buf = ""
    mode = "plain"
    while i < len(line):
        if mode == "plain" and line.startswith("**", i):
            if buf:
                paragraph.add_run(buf)
                buf = ""
            mode = "bold"
            i += 2
            continue
        if mode == "bold" and line.startswith("**", i):
            run = paragraph.add_run(buf)
            run.bold = True
            buf = ""
            mode = "plain"
            i += 2
            continue
        if mode == "plain" and line.startswith("`", i):
            if buf:
                paragraph.add_run(buf)
                buf = ""
            mode = "code"
            i += 1
            continue
        if mode == "code" and line.startswith("`", i):
            run = paragraph.add_run(buf)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            buf = ""
            mode = "plain"
            i += 1
            continue
        buf += line[i]
        i += 1
    if buf:
        run = paragraph.add_run(buf)
        if mode == "bold":
            run.bold = True
        if mode == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)


def is_separator_row(cells: list[str]) -> bool:
    for c in cells:
        cleaned = c.replace(":", "").replace("-", "").strip()
        if cleaned:
            return False
    return True


def flush_table(doc: Document, table_rows: list[str]) -> None:
    if not table_rows:
        return
    rows: list[list[str]] = []
    for r in table_rows:
        cells = [c.strip() for c in r.strip("|").split("|")]
        if is_separator_row(cells):
            continue
        rows.append(cells)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs_with_code(p, cell_text)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph("")


def main() -> None:
    text = MD.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    in_code = False
    code_lines: list[str] = []
    table_rows: list[str] = []

    for raw in text.splitlines():
        line = raw.rstrip("\n")
        if line.startswith("```"):
            if not in_code:
                flush_table(doc, table_rows)
                table_rows = []
                in_code = True
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(8)
                code_lines = []
            continue
        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("|") and "|" in line[1:]:
            table_rows.append(line)
            continue

        flush_table(doc, table_rows)
        table_rows = []

        if not line.strip() or line.strip() == "---":
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            continue
        if line.lstrip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_code(p, line.lstrip()[2:])
            continue

        stripped = line.lstrip()
        if stripped and stripped[0].isdigit() and ". " in stripped[:4]:
            content = stripped.split(". ", 1)[1]
            p = doc.add_paragraph(style="List Number")
            add_runs_with_code(p, content)
            continue

        p = doc.add_paragraph()
        add_runs_with_code(p, line)

    flush_table(doc, table_rows)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Size_bytes {OUT.stat().st_size}")


if __name__ == "__main__":
    main()
